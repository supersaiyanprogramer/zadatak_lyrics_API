
import requests
from docx import Document

#pitaj korisnika za unos izvođača i pjesme
artist = input("unesite naziv izvođača: ")
title = input("unesite naziv pjesme: ")


#ime api-ja
#originalni link - https://api.lyrics.ovh/v1/artist/title
api_url="https://api.lyrics.ovh/v1/{}/{}.format(artist, title)"


#dohvati podatke s apija, nazovi te podatke feedback, dekodiraj ih u utf-8 u slučaju Č,Ć,Ž,Đ itd.
response = requests.get(api_url)
feedback = response.read()
feedback = feedback('utf-8')

#printaj feedback, odnosno tekst pjesme
print(feedback)

response.close()


#podatke nazovi data i otvori ih u wordu s odgovarajućim parametrima
data = json.loads(text)

#otvori word dokument s headingom i paragrafom
document = Document()

document.add_heading('Naziv izvođača i naziv pjesme', 0)
p = document.add_paragraph(feedback)


document.add_page_break()

#spremi word dokument sa zadanim imenom
document.save('artist_song.docx')
